"""
Shared pytest fixtures for all phases.
"""

import io
from unittest.mock import MagicMock, patch

import pytest
import chromadb

import app.services.ingestion as ingestion_module


# ── SQLite DB isolation ───────────────────────────────────────────────────── #

@pytest.fixture(autouse=True)
def reset_db(tmp_path):
    """Give every test its own empty SQLite DB."""
    import app.core.database as db_module
    original = db_module.DB_PATH
    db_module.DB_PATH = str(tmp_path / "test.db")
    db_module.init_db()
    yield str(tmp_path / "test.db")
    db_module.DB_PATH = original


# ── Rate limiter isolation ─────────────────────────────────────────────────── #

@pytest.fixture(autouse=True)
def reset_rate_limiter():
    """Reset rate limiter state between tests."""
    from app.core.rate_limiter import reset
    reset()
    yield
    reset()


# ── ChromaDB isolation ────────────────────────────────────────────────────── #

@pytest.fixture(autouse=True)
def reset_chroma():
    """Give every test its own empty in-memory ChromaDB."""
    ingestion_module._chroma_client = chromadb.EphemeralClient()
    yield
    ingestion_module._chroma_client = None


# ── Embedding mock (avoids real Gemini API calls in all tests) ────────────── #

@pytest.fixture(autouse=True)
def mock_embed_texts():
    """
    Replace embed_texts with a deterministic fake so no GOOGLE_API_KEY is
    needed in CI.  All texts get the same 8-dim unit vector; cosine similarity
    between any two is 1.0 (score = 1.0), which satisfies the >0.5 threshold.
    """
    _DIM = 8
    _VEC = [1.0] + [0.0] * (_DIM - 1)  # unit vector along first axis

    def _fake(texts: list[str]) -> list[list[float]]:
        return [_VEC for _ in texts]

    with patch("app.services.ingestion.embed_texts", side_effect=_fake), \
         patch("app.services.retrieval.embed_texts", side_effect=_fake), \
         patch("app.services.rag_chain.embed_texts", side_effect=_fake):
        yield


# ── LLM mock (avoids real Gemini calls in unit tests) ─────────────────────── #

@pytest.fixture
def mock_llm():
    """Patch ChatGoogleGenerativeAI so tests never hit the Gemini API."""
    def _fake_invoke(messages):
        prompt = messages[0].content if messages else ""
        if "GDP of Mars" in prompt or "mars" in prompt.lower():
            text = "I don't have enough information in the provided documents to answer this question."
        else:
            text = "The document discusses qualitative and quantitative methodology frameworks used in social science research."
        resp = MagicMock()
        resp.content = text
        return resp

    with patch("app.services.rag_chain.ChatGoogleGenerativeAI") as mock_cls:
        instance = MagicMock()
        instance.invoke.side_effect = _fake_invoke
        mock_cls.return_value = instance
        yield instance


# ── Document fixtures ─────────────────────────────────────────────────────── #

def _build_minimal_pdf(text: str) -> bytes:
    """
    Build a minimal but valid PDF-1.4 with one page containing *text*.
    Parentheses and backslashes in *text* must be escaped for PDF string syntax.
    """
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content_stream = f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET".encode()

    objs: list[bytes] = [
        b"1 0 obj\n<</Type /Catalog /Pages 2 0 R>>\nendobj\n",
        b"2 0 obj\n<</Type /Pages /Kids [3 0 R] /Count 1>>\nendobj\n",
        (
            b"3 0 obj\n<</Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
            b" /Contents 4 0 R /Resources <</Font <</F1 5 0 R>>>>>>\nendobj\n"
        ),
        f"4 0 obj\n<</Length {len(content_stream)}>>\nstream\n".encode()
        + content_stream
        + b"\nendstream\nendobj\n",
        b"5 0 obj\n<</Type /Font /Subtype /Type1 /BaseFont /Helvetica>>\nendobj\n",
    ]

    pdf = b"%PDF-1.4\n"
    offsets: list[int] = []
    for obj in objs:
        offsets.append(len(pdf))
        pdf += obj

    xref_pos = len(pdf)
    xref = f"xref\n0 {len(objs) + 1}\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"
    trailer = (
        f"trailer\n<</Size {len(objs) + 1} /Root 1 0 R>>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    )
    pdf += (xref + trailer).encode()
    return pdf


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    return _build_minimal_pdf(
        "This document discusses methodology and research approaches in social science."
    )


@pytest.fixture
def sample_docx_bytes() -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Research Methodology", 0)
    doc.add_paragraph(
        "This document discusses methodology and research approaches. "
        "The methodology section outlines the key research methods used in social science."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_txt_bytes() -> bytes:
    return (
        b"This document discusses methodology and research approaches in social science. "
        b"The study examines various methodological frameworks used in field research. "
        b"Qualitative and quantitative methods are both covered extensively."
    )
